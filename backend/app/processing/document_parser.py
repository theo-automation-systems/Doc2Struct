"""
Document parsing module — supports PDF, DOCX, XLSX, and plain text.
Extracts raw text and metadata from uploaded files.
"""

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse documents of various formats and extract raw text content."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".doc", ".xlsx", ".xls"}

    def parse(self, file_bytes: bytes, filename: str) -> dict:
        """
        Parse a document and return structured text content.

        Returns:
            {
                "text": str,
                "pages": int,
                "metadata": dict,
                "tables": list,
                "error": str | None
            }
        """
        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return {"text": "", "pages": 0, "metadata": {}, "tables": [], "error": f"Unsupported file type: {ext}"}

        try:
            if ext == ".pdf":
                return self._parse_pdf(file_bytes)
            elif ext == ".txt":
                return self._parse_text(file_bytes)
            elif ext in (".docx", ".doc"):
                return self._parse_docx(file_bytes)
            elif ext in (".xlsx", ".xls"):
                return self._parse_excel(file_bytes)
        except Exception as e:
            logger.exception(f"Failed to parse {filename}")
            return {"text": "", "pages": 0, "metadata": {}, "tables": [], "error": str(e)}

    def _parse_pdf(self, file_bytes: bytes) -> dict:
        try:
            import pdfplumber
            text_parts = []
            tables = []
            page_count = 0

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)

            return {
                "text": "\n\n".join(text_parts),
                "pages": page_count,
                "metadata": {"parser": "pdfplumber"},
                "tables": tables,
                "error": None,
            }
        except ImportError:
            # Fallback to PyMuPDF
            return self._parse_pdf_fitz(file_bytes)

    def _parse_pdf_fitz(self, file_bytes: bytes) -> dict:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []

        for page in doc:
            text_parts.append(page.get_text())

        return {
            "text": "\n\n".join(text_parts),
            "pages": len(doc),
            "metadata": {"parser": "pymupdf"},
            "tables": [],
            "error": None,
        }

    def _parse_text(self, file_bytes: bytes) -> dict:
        text = file_bytes.decode("utf-8", errors="replace")
        lines = text.splitlines()
        return {
            "text": text,
            "pages": max(1, len(lines) // 50),
            "metadata": {"parser": "plaintext"},
            "tables": [],
            "error": None,
        }

    def _parse_docx(self, file_bytes: bytes) -> dict:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []

        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(rows)

        return {
            "text": "\n".join(paragraphs),
            "pages": 1,
            "metadata": {"parser": "python-docx"},
            "tables": tables,
            "error": None,
        }

    def _parse_excel(self, file_bytes: bytes) -> dict:
        import pandas as pd
        xl = pd.ExcelFile(io.BytesIO(file_bytes))
        all_text = []
        all_tables = []

        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            all_text.append(f"=== Sheet: {sheet_name} ===\n{df.to_string()}")
            all_tables.append(df.values.tolist())

        return {
            "text": "\n\n".join(all_text),
            "pages": len(xl.sheet_names),
            "metadata": {"parser": "pandas", "sheets": xl.sheet_names},
            "tables": all_tables,
            "error": None,
        }
